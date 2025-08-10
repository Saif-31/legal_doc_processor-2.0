# legal_doc_processor.py
import os
import re
import logging
from datetime import datetime
import io
import copy
from typing import Iterator, List, Optional, Union, Dict, Tuple
import streamlit as st
from docx import Document
from docx.document import Document as DocumentObject
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph, Run
from docx.table import Table, _Row
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl, CT_Row
from difflib import SequenceMatcher
from dotenv import load_dotenv

# --- Configuration and Setup ---
load_dotenv()
LOG_DIR = "logs"
os.makedirs(LOG_DIR, exist_ok=True)

# Setup logging
log_file = os.path.join(LOG_DIR, f"processor_{datetime.now().strftime('%Y%m%d')}.log")

# Open in write mode to clear it if it exists, without deleting
with open(log_file, 'w'):
    pass  # just clears contents

logging.basicConfig(
    filename=log_file,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)


logging.basicConfig(
    filename=log_file,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

Block = Union[Paragraph, Table]
DiffableBlock = Union[Paragraph, _Row]

# --- Constants ---
COLOR_GREEN_RGB = RGBColor(0, 204, 51)
COLOR_RED_RGB = RGBColor(255, 0, 0)
ARTICLE_RE = re.compile(r"^\s*Član\s+(\d+[a-zA-Z]*)", re.IGNORECASE)

# --- High-Fidelity Deep Copy Functions ---
def deep_copy_paragraph(source_p: Paragraph, target_doc: DocumentObject) -> Paragraph:
    new_p_element = copy.deepcopy(source_p._element)
    target_doc.element.body.append(new_p_element)
    return Paragraph(new_p_element, target_doc)

def deep_copy_row(source_row: _Row, target_table: Table) -> _Row:
    """Deep copies a table row into a target table."""
    new_row_element = copy.deepcopy(source_row._element)
    target_table._tbl.append(new_row_element)
    return _Row(new_row_element, target_table)

# --- Coloring and Formatting Helpers ---
def iter_block_items(doc: DocumentObject) -> Iterator[Block]:
    """Iterates over paragraphs and tables in a document in their original order."""
    for child in doc.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)
            
def color_paragraph_runs(p: Paragraph, color: RGBColor):
    for run in p.runs:
        run.font.color.rgb = color

def add_brackets_to_paragraph(p: Paragraph, color: RGBColor):
    """Adds brackets around the entire paragraph text with proper coloring."""
    # Get the original text
    original_text = p.text
    
    # Clear the paragraph
    p.clear()
    
    # Add opening bracket
    p.add_run('[').font.name = 'Arial'
    
    # Add the colored text
    colored_run = p.add_run(original_text)
    colored_run.font.color.rgb = color
    colored_run.font.name = 'Arial'
    
    # Add closing bracket
    p.add_run(']').font.name = 'Arial'

def color_row_runs(row: _Row, color: RGBColor):
    for cell in row.cells:
        for p in cell.paragraphs:
            color_paragraph_runs(p, color)

# --- Part B: DEFINITIVE IMPLEMENTATION ---

def add_explanatory_and_title_tables(diff_doc: DocumentObject, new_doc: DocumentObject):
    """
    Adds the two standard explanation headers with correct formatting.
    """
    # 1. Add the grey explanation box with full text
    table1 = diff_doc.add_table(rows=1, cols=1)
    cell1 = table1.rows[0].cells[0]
    tcPr1 = cell1._tc.get_or_add_tcPr()
    shd1 = OxmlElement('w:shd')
    shd1.set(qn('w:val'), 'clear')
    shd1.set(qn('w:fill'), 'CCCCCC')
    tcPr1.append(shd1)
    
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.clear()
    
    # Add the full explanation text with proper formatting
    p1.add_run("Radi lakšeg sagledavanja izmena i dopuna propisa, nova sadržina odredaba data je ").bold = True
    r_green = p1.add_run("zelenom")
    r_green.bold = True
    shd_green = OxmlElement('w:shd')
    shd_green.set(qn('w:val'), 'clear')
    shd_green.set(qn('w:fill'), '33FF33')
    r_green.element.get_or_add_rPr().append(shd_green)
    
    p1.add_run(", prethodna ").bold = True
    r_red = p1.add_run("crvenom")
    r_red.bold = True
    r_red.font.color.rgb = COLOR_RED_RGB
    
    p1.add_run(" bojom, a nepromenjene odredbe nisu posebno označene, tako da pregledanjem crno-zelenog teksta pregledate važeću, a crno-crvenog teksta, prethodnu verziju propisa. Prečišćen tekst bez crvenih i zelenih oznaka i dalje možete videti na tabu ").bold = True
    
    r_tab = p1.add_run('"Tekst dokumenta".')
    r_tab.bold = True
    r_tab.italic = True
    
    for run in p1.runs:
        run.font.name = 'Arial'

    # 2. Add the combined title table with correct colors
    table2 = diff_doc.add_table(rows=1, cols=1)
    cell2 = table2.rows[0].cells[0]
    tcPr2 = cell2._tc.get_or_add_tcPr()
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'), 'clear')
    shd2.set(qn('w:fill'), '8A084B')  # Burgundy background
    tcPr2.append(shd2)
    
    # First paragraph - analytical view with black background
    p2_1 = cell2.paragraphs[0]
    p2_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2_1.clear()
    
    # Add black shading to the paragraph
    pPr = p2_1._element.get_or_add_pPr()
    shd_p = OxmlElement('w:shd')
    shd_p.set(qn('w:val'), 'clear')
    shd_p.set(qn('w:fill'), '000000')
    pPr.append(shd_p)
    
    run2_1 = p2_1.add_run("Propis - analitički prikaz promena")
    run2_1.bold = True
    run2_1.italic = True
    run2_1.font.name = 'Arial'
    run2_1.font.size = Pt(13)
    run2_1.font.color.rgb = RGBColor(255, 232, 191)
    
    # Extract title information from the new document
    title_info = extract_title_info(new_doc)
    
    # Add ZAKON paragraph
    p2_2 = cell2.add_paragraph()
    p2_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2_2 = p2_2.add_run(title_info['law_type'])
    run2_2.bold = True
    run2_2.font.name = 'Arial'
    run2_2.font.size = Pt(18)
    run2_2.font.color.rgb = RGBColor(255, 232, 191)
    
    # Add law name paragraph
    p2_3 = cell2.add_paragraph()
    p2_3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2_3 = p2_3.add_run(title_info['law_name'])
    run2_3.bold = True
    run2_3.font.name = 'Arial'
    run2_3.font.size = Pt(17)
    run2_3.font.color.rgb = RGBColor(255, 255, 255)
    
    # Add gazette info with black background
    p2_4 = cell2.add_paragraph()
    p2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add black shading to the gazette paragraph
    pPr4 = p2_4._element.get_or_add_pPr()
    shd_p4 = OxmlElement('w:shd')
    shd_p4.set(qn('w:val'), 'clear')
    shd_p4.set(qn('w:fill'), '000000')
    pPr4.append(shd_p4)
    
    run2_4 = p2_4.add_run(title_info['gazette'])
    run2_4.italic = True
    run2_4.font.name = 'Arial'
    run2_4.font.size = Pt(13)
    run2_4.font.color.rgb = RGBColor(255, 232, 191)

def extract_title_info(doc: DocumentObject) -> Dict[str, str]:
    """Extract title information from the document."""
    title_info = {
        'law_type': 'ZAKON',
        'law_name': 'O RAČUNOVODSTVU',
        'gazette': '("Sl. glasnik RS", br. 73/2019 i 44/2021 - dr. zakon)'
    }
    
    # Try to extract from the first table if it exists
    first_block = next(iter_block_items(doc), None)
    if isinstance(first_block, Table) and first_block.rows:
        cell_text = first_block.rows[0].cells[0].text
        lines = cell_text.strip().split('\n')
        
        for i, line in enumerate(lines):
            if 'ZAKON' in line.upper() and i == 0:
                title_info['law_type'] = line.strip()
            elif i == 1 and 'O ' in line.upper():
                title_info['law_name'] = line.strip()
            elif 'Sl. glasnik' in line:
                title_info['gazette'] = line.strip()
    
    return title_info

def flatten_document_for_diff(doc: DocumentObject) -> List[DiffableBlock]:
    """Flattens a document into a list of Paragraphs and Table Rows for granular diffing."""
    flat_list = []
    # Skip the first block if it's a table (the main title header)
    all_blocks = list(iter_block_items(doc))
    start_index = 1 if all_blocks and isinstance(all_blocks[0], Table) else 0
    
    for block in all_blocks[start_index:]:
        if isinstance(block, Paragraph):
            flat_list.append(block)
        elif isinstance(block, Table):
            for row in block.rows:
                flat_list.append(row)
    return flat_list

def get_text_from_diffable_block(block: DiffableBlock) -> str:
    """Gets a consistent, normalized text string from a Paragraph or a Table Row."""
    text = ""
    if isinstance(block, _Row):
        text = " | ".join(cell.text for cell in block.cells)
    elif isinstance(block, Paragraph):
        text = block.text
    # Normalize: remove asterisks, non-breaking spaces, and leading/trailing whitespace
    return text.replace('*', '').replace('\u00A0', ' ').strip()

def extract_amending_ref(gov_doc: DocumentObject) -> List[str]:
    """
    Extracts a list of amending references from the government document.
    One reference per 'Član X' paragraph, using law name and gazette from the first table.
    Returns: ["[ČLAN 5 ODLUKA O... (\"SL. GLASNIK RS\", BR. ...)]", ...]
    """
    law_name = ""
    gazette = ""
    
    # 1. Extract law name and gazette from the first table
    for block in iter_block_items(gov_doc):
        if isinstance(block, Table):
            texts = [
                p.text.strip()
                for row in block.rows
                for cell in row.cells
                for p in cell.paragraphs
                if p.text.strip()
            ]
            for t in texts:
                if not law_name and re.search(r"zakon\s+o\s+", t, re.IGNORECASE):
                    law_name = t.upper()
                if not gazette and re.search(r"glasnik\s+rs", t, re.IGNORECASE):
                    gazette = t.upper()
            break  # Only first table

    # 2. Generate one reference per 'Član ...' paragraph
    references = []
    for para in gov_doc.paragraphs:
        text = para.text.strip()
        if re.match(r"^Član\s+\d+(\s+stav\s+\d+)?", text, re.IGNORECASE):
            article = text.upper()
            ref = f"[{article} {law_name} {gazette}]"
            references.append(ref)

    if not references:
        logging.warning("No amending references extracted. Using fallback.")
        return ["[ČLAN 23 STAV 2 ZAKON O ELEKTRONSKOM FAKE FAKTURISANJU (\"SL. GLASNIK RS\", BR. 44/2021)]"]

    return references

def insert_amending_references(doc: DocumentObject, gov_doc: DocumentObject):
    """
    Inserts a green reference before each modified article.
    Each reference is based on the article title and data from the amendment doc.
    Uses your working logic: one ref per Član, from first table.
    """
    # Get all references (one per Član in the amendment doc)
    refs = extract_amending_ref(gov_doc)
    if not refs:
        logging.warning("No references to insert.")
        return

    # We’ll use the same set of references for all modified articles
    # (If you want 1:1 mapping, we can improve this later)
    ref_text = refs[0]  # Use first reference (all are usually similar)
    ref_text_clean = ref_text.strip("[]")  # Remove brackets for formatting

    # === Find all article titles that have changes ===
    articles_to_mark = set()
    current_article_title_p = None

    for block in iter_block_items(doc):
        # Detect article title: "Član X", ignore '*'
        if isinstance(block, Paragraph):
            clean_text = block.text.replace('*', '').strip()
            if ARTICLE_RE.match(clean_text):
                current_article_title_p = block

        # Check if this block has changes (red or green)
        has_change = False
        if isinstance(block, Paragraph):
            has_change = any(
                run.font.color.rgb in [COLOR_GREEN_RGB, COLOR_RED_RGB]
                for run in block.runs
            )
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if any(
                            run.font.color.rgb in [COLOR_GREEN_RGB, COLOR_RED_RGB]
                            for run in p.runs
                        ):
                            has_change = True
                            break
                    if has_change:
                        break
                if has_change:
                    break

        if has_change and current_article_title_p:
            articles_to_mark.add(current_article_title_p)

    # === Insert reference before each marked article ===
    for p_title in sorted(
        articles_to_mark,
        key=lambda p: doc.paragraphs.index(p) if p in doc.paragraphs else float('inf')
    ):
        # Insert new paragraph before the article
        ref_p = p_title.insert_paragraph_before()
        ref_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_p.paragraph_format.space_before = Pt(12)
        ref_p.paragraph_format.space_after = Pt(6)
        ref_p.paragraph_format.line_spacing = 1.0

        # Add formatted reference: [ ... ]
        ref_p.add_run('[').bold = True
        run = ref_p.add_run(ref_text_clean)
        run.bold = True
        run.font.color.rgb = COLOR_GREEN_RGB
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        ref_p.add_run(']').bold = True

        # Ensure all runs use Arial 12pt
        for run in ref_p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)

    logging.info(f"Inserted amending references before {len(articles_to_mark)} articles.")
  
def process_part_b_streamlit(orig_file, new_file, amend_file) -> DocumentObject:
    """
    Processes Part B with cell-level table diffing and correct formatting.
    """
    try:
        orig_doc = Document(orig_file)
        new_doc = Document(new_file)
        gov_doc = Document(amend_file)

        # 1. Create a pristine, blank document based on the original's template
        diff_doc = Document(orig_file)
        diff_doc.element.body.clear_content()
        
        # 2. Add the custom headers and the final title table
        add_explanatory_and_title_tables(diff_doc, new_doc)
        
        # 3. Flatten documents into lists of paragraphs and table rows
        orig_flat_list = flatten_document_for_diff(orig_doc)
        new_flat_list = flatten_document_for_diff(new_doc)
        orig_texts = [get_text_from_diffable_block(b) for b in orig_flat_list]
        new_texts = [get_text_from_diffable_block(b) for b in new_flat_list]

        # Create lookup maps from XML element to Table object
        orig_table_map = {tbl._element: tbl for tbl in orig_doc.tables}
        new_table_map = {tbl._element: tbl for tbl in new_doc.tables}

        # 4. Perform granular diff and build the document
        current_table_in_diff = None
        for tag, i1, i2, j1, j2 in SequenceMatcher(None, orig_texts, new_texts, autojunk=False).get_opcodes():
            
            # Helper to create a new table in the diff doc if we encounter a row
            def ensure_table_context(source_row: _Row, table_map: dict) -> Table:
                nonlocal current_table_in_diff
                # Get the parent _tbl XML element of the row's _tr element
                source_tbl_element = source_row._tr.getparent()
                # Use the map to find the corresponding high-level Table object
                source_table_object = table_map[source_tbl_element]
                
                # If we're starting a new table in the diff doc...
                if current_table_in_diff is None or current_table_in_diff.style != source_table_object.style:
                    # Create a new, empty table shell in the diff doc
                    new_table_element = copy.deepcopy(source_table_object._element)
                    
                    # Correctly remove all rows from the XML element
                    for tr in new_table_element.xpath('./w:tr'):
                        tr.getparent().remove(tr)
                        
                    diff_doc.element.body.append(new_table_element)
                    current_table_in_diff = Table(new_table_element, diff_doc)
                return current_table_in_diff
            
            # Reset table context if the previous block was not a row
            prev_block = orig_flat_list[i1-1] if i1 > 0 else None
            if not isinstance(prev_block, _Row):
                current_table_in_diff = None

            if tag == 'equal':
                for k in range(j1, j2):
                    block = new_flat_list[k]
                    if isinstance(block, Paragraph):
                        current_table_in_diff = None
                        # Handle article titles without asterisks
                        new_p = deep_copy_paragraph(block, diff_doc)
                        # If it's an article title with asterisk, remove it
                        if ARTICLE_RE.match(block.text) and '*' in block.text:
                            new_p.clear()
                            for run in block.runs:
                                new_text = run.text.replace('*', '')
                                if new_text:
                                    new_run = new_p.add_run(new_text)
                                    # Copy formatting
                                    if run.bold: new_run.bold = True
                                    if run.italic: new_run.italic = True
                                    if run.font.name: new_run.font.name = run.font.name
                                    if run.font.size: new_run.font.size = run.font.size
                    elif isinstance(block, _Row):
                        target_table = ensure_table_context(block, new_table_map)
                        deep_copy_row(block, target_table)

            elif tag == 'delete' or tag == 'replace':
                for k in range(i1, i2):
                    block = orig_flat_list[k]
                    if isinstance(block, Paragraph):
                        current_table_in_diff = None
                        p = deep_copy_paragraph(block, diff_doc)
                        # Handle article titles
                        if ARTICLE_RE.match(p.text):
                            # Remove asterisk if present
                            if '*' in p.text:
                                clean_text = p.text.replace('*', '')
                                p.clear()
                                run = p.add_run(clean_text)
                                run.bold = True
                                run.font.name = 'Arial'
                                run.font.size = Pt(12)
                        else:
                            # For non-article paragraphs, add brackets properly
                            add_brackets_to_paragraph(p, COLOR_RED_RGB)
                    elif isinstance(block, _Row):
                        target_table = ensure_table_context(block, orig_table_map)
                        new_row = deep_copy_row(block, target_table)
                        color_row_runs(new_row, COLOR_RED_RGB)

            if tag == 'insert' or tag == 'replace':
                for k in range(j1, j2):
                    block = new_flat_list[k]
                    if isinstance(block, Paragraph):
                        current_table_in_diff = None
                        p = deep_copy_paragraph(block, diff_doc)
                        # Handle article titles
                        if ARTICLE_RE.match(p.text):
                            # Remove asterisk if present
                            if '*' in p.text:
                                clean_text = p.text.replace('*', '')
                                p.clear()
                                run = p.add_run(clean_text)
                                run.bold = True
                                run.font.name = 'Arial'
                                run.font.size = Pt(12)
                        else:
                            # For non-article paragraphs, add brackets properly
                            add_brackets_to_paragraph(p, COLOR_GREEN_RGB)
                    elif isinstance(block, _Row):
                        target_table = ensure_table_context(block, new_table_map)
                        new_row = deep_copy_row(block, target_table)
                        color_row_runs(new_row, COLOR_GREEN_RGB)

        # 5. Post-processing
        insert_amending_references(diff_doc, gov_doc)
        
        return diff_doc
    except Exception as e:
        logging.error(f"Error in process_part_b_streamlit: {e}", exc_info=True)
        raise

# --- Streamlit App ---
def main():
    st.set_page_config(layout="wide", page_title="Professional Legal Document Processor")
    st.title("Professional Legal Document Processor")
    st.markdown("This tool creates high-fidelity 'redline' comparisons of legal documents, preserving all formatting.")
    st.markdown("---")

    st.header("1. Upload Files for Processing")
    st.info("For the Colored Diff, please provide all three documents.")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        orig_file = st.file_uploader("A. Original Law (`.docx`)", type="docx", key="orig")
    with col2:
        new_file = st.file_uploader("B. New Consolidated Law (`.docx`)", type="docx", key="new")
    with col3:
        amend_file = st.file_uploader("C. Government Changes (`.docx`)", type="docx", key="amend")
    
    st.markdown("---")

    # --- Processing Section ---
    st.header("2. Generate Document")
    if st.button("Generate Colored Diff (Part B)", type="primary", use_container_width=True):
        if not orig_file or not new_file or not amend_file:
            st.error("Please upload all three files to generate the Colored Diff.")
        else:
            with st.spinner("Processing Part B: Generating high-fidelity colored diff..."):
                try:
                    diff_doc = process_part_b_streamlit(orig_file, new_file, amend_file)
                    st.session_state['diff_doc'] = diff_doc
                    st.success("Part B completed successfully! Your document is ready for download.")
                    logging.info("Part B successfully completed.")
                except Exception as e:
                     st.error(f"An error occurred during Part B processing: {e}")
                     logging.error(f"Part B error: {e}", exc_info=True)

    st.markdown("---")

    # --- Download Section ---
    st.header("3. Download Result")
    if 'diff_doc' in st.session_state:
        # Save the generated document to a memory buffer
        diff_buffer = io.BytesIO()
        st.session_state['diff_doc'].save(diff_buffer)
        diff_buffer.seek(0)

        # Create the download button
        st.download_button(
            label="Download Colored_Diff.docx",
            data=diff_buffer,
            file_name="Colored_Diff.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.info("Process a document to enable the download button.")

if __name__ == "__main__":
    main()